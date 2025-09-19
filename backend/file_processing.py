import os
import tempfile
import textwrap
from pathlib import Path
from fastapi import HTTPException, UploadFile
from fastapi.responses import FileResponse
import fitz  # PyMuPDF for PDF processing
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import io
from config import config
from models import FileValidationResult

def validate_uploaded_file(file: UploadFile) -> FileValidationResult:
    """Comprehensive file validation"""
    # Check file size (estimate from filename and content type)
    file_size = 0
    if hasattr(file, 'size') and file.size:
        file_size = file.size

    # Validate file type
    if file.content_type not in config.SUPPORTED_FILE_TYPES:
        return FileValidationResult(
            is_valid=False,
            error_message=f"Unsupported file type: {file.content_type}. Only PDF and DOCX files are supported.",
            file_size=file_size,
            file_type=file.content_type or "unknown"
        )

    # Validate file extension
    if file.filename:
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in config.SUPPORTED_FILE_TYPES[file.content_type]:
            return FileValidationResult(
                is_valid=False,
                error_message=f"File extension {file_extension} doesn't match content type {file.content_type}",
                file_size=file_size,
                file_type=file.content_type
            )

    return FileValidationResult(
        is_valid=True,
        file_size=file_size,
        file_type=file.content_type
    )

def extract_text_from_pdf(file_content: bytes) -> dict:
    """Enhanced PDF text extraction with metadata"""
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        text_content = ""
        page_count = len(doc)

        # Extract text from all pages
        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                text_content += f"--- Page {page_num + 1} ---\n{page_text}\n\n"

        doc.close()

        # Validate extracted text
        if not text_content.strip():
            raise ValueError("No readable text found in the PDF. The file might be image-based or corrupted.")

        # Calculate statistics
        word_count = len(text_content.split())
        character_count = len(text_content)

        return {
            "text": text_content.strip(),
            "page_count": page_count,
            "word_count": word_count,
            "character_count": character_count,
            "extraction_method": "PyMuPDF"
        }

    except Exception as e:
        if "No readable text found" in str(e):
            raise HTTPException(status_code=400, detail=str(e))
        else:
            raise HTTPException(status_code=400, detail=f"Error processing PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> dict:
    """Enhanced DOCX text extraction with metadata"""
    try:
        doc = Document(io.BytesIO(file_content))
        text_content = ""
        paragraph_count = 0

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
                paragraph_count += 1

        # Extract text from tables if any
        table_content = ""
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_content += " | ".join(row_text) + "\n"

        if table_content:
            text_content += "\n--- Tables ---\n" + table_content

        # Validate extracted text
        if not text_content.strip():
            raise ValueError("No readable text found in the DOCX file. The file might be empty or corrupted.")

        # Calculate statistics
        word_count = len(text_content.split())
        character_count = len(text_content)

        return {
            "text": text_content.strip(),
            "paragraph_count": paragraph_count,
            "table_count": len(doc.tables),
            "word_count": word_count,
            "character_count": character_count,
            "extraction_method": "python-docx"
        }

    except Exception as e:
        if "No readable text found" in str(e):
            raise HTTPException(status_code=400, detail=str(e))
        else:
            raise HTTPException(status_code=400, detail=f"Error processing DOCX: {str(e)}")

def create_pdf_from_text(text: str, filename: str) -> str:
    """Create PDF from cleaned text"""
    temp_dir = tempfile.mkdtemp()
    pdf_path = os.path.join(temp_dir, f"{filename}.pdf")

    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()

    # Create a custom style for resume content
    resume_style = ParagraphStyle(
        'ResumeStyle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        leftIndent=0,
        rightIndent=0,
    )

    story = []

    # Split text into paragraphs and create PDF content
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            # Wrap long lines
            wrapped_text = textwrap.fill(para, width=80)
            story.append(Paragraph(wrapped_text, resume_style))
        else:
            story.append(Spacer(1, 12))

    doc.build(story)
    return pdf_path

def extract_text_from_file(file_content: bytes, content_type: str) -> dict:
    """Extract text from file based on content type"""
    if content_type == "application/pdf":
        return extract_text_from_pdf(file_content)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file_content)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {content_type}")