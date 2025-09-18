from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Depends, Form
from fastapi.responses import FileResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional
import uuid
from datetime import datetime, timezone
import asyncio
import tempfile
import shutil
from passlib.context import CryptContext
from jose import JWTError, jwt
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import io
import fitz  # PyMuPDF for PDF processing
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import textwrap

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT and password handling
SECRET_KEY = os.environ.get("SECRET_KEY", "your-secret-key-here")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
security = HTTPBearer()

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# AI Integration
from emergentintegrations.llm.chat import LlmChat, UserMessage

# Initialize AI chat
async def get_ai_chat():
    chat = LlmChat(
        api_key=os.environ.get("EMERGENT_LLM_KEY"),
        session_id=str(uuid.uuid4()),
        system_message="You are a professional resume editor. Clean the following resume text by fixing all grammar and punctuation errors. Do not change the meaning, content structure, or formatting. Only fix grammatical errors, spelling mistakes, and punctuation. Return the corrected text maintaining the original structure and format."
    ).with_model("gemini", "gemini-2.5-pro")
    return chat

# Define Models
class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    username: str
    email: str
    hashed_password: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class UserCreate(BaseModel):
    username: str
    email: str
    password: str

class UserLogin(BaseModel):
    email: str
    password: str

class Token(BaseModel):
    access_token: str
    token_type: str

class Resume(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    original_filename: str
    original_text: str
    cleaned_text: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ResumeCleanRequest(BaseModel):
    resume_id: str

# Helper functions
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.now(timezone.utc).timestamp() + (ACCESS_TOKEN_EXPIRE_MINUTES * 60)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise HTTPException(status_code=401, detail="Invalid authentication credentials")
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid authentication credentials")
    
    user = await db.users.find_one({"email": email})
    if user is None:
        raise HTTPException(status_code=401, detail="User not found")
    
    return User(**user)

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing DOCX: {str(e)}")

def create_pdf_from_text(text: str, filename: str) -> str:
    """Create PDF from cleaned text"""
    temp_dir = tempfile.mkdtemp()
    pdf_path = os.path.join(temp_dir, f"{filename}.pdf")
    
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create a custom style for resume content
    resume_style = ParagraphStyle(
        'ResumeStyle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        leftIndent=0,
        rightIndent=0,
    )
    
    story = []
    
    # Split text into paragraphs and create PDF content
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            # Wrap long lines
            wrapped_text = textwrap.fill(para, width=80)
            story.append(Paragraph(wrapped_text, resume_style))
        else:
            story.append(Spacer(1, 12))
    
    doc.build(story)
    return pdf_path

# Authentication routes
@api_router.post("/auth/register", response_model=Token)
async def register(user_data: UserCreate):
    # Check if user already exists
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create new user
    hashed_password = get_password_hash(user_data.password)
    user = User(
        username=user_data.username,
        email=user_data.email,
        hashed_password=hashed_password
    )
    
    await db.users.insert_one(user.dict())
    
    # Create access token
    access_token = create_access_token(data={"sub": user.email})
    return {"access_token": access_token, "token_type": "bearer"}

@api_router.post("/auth/login", response_model=Token)
async def login(user_data: UserLogin):
    user = await db.users.find_one({"email": user_data.email})
    if not user or not verify_password(user_data.password, user["hashed_password"]):
        raise HTTPException(status_code=401, detail="Incorrect email or password")
    
    access_token = create_access_token(data={"sub": user["email"]})
    return {"access_token": access_token, "token_type": "bearer"}

@api_router.get("/auth/me", response_model=User)
async def get_me(current_user: User = Depends(get_current_user)):
    return current_user

# Enhanced Models for file upload
class FileUploadResponse(BaseModel):
    message: str
    resume_id: str
    original_text: str
    file_size: int
    word_count: int
    character_count: int
    file_type: str

class FileValidationResult(BaseModel):
    is_valid: bool
    error_message: Optional[str] = None
    file_size: int
    file_type: str

# Enhanced file validation function
def validate_uploaded_file(file: UploadFile) -> FileValidationResult:
    """Comprehensive file validation"""
    # Maximum file size: 10MB
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB in bytes
    
    # Supported MIME types and extensions
    SUPPORTED_TYPES = {
        "application/pdf": [".pdf"],
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": [".docx"]
    }
    
    # Check file size (estimate from filename and content type)
    file_size = 0
    if hasattr(file, 'size') and file.size:
        file_size = file.size
    
    # Validate file type
    if file.content_type not in SUPPORTED_TYPES:
        return FileValidationResult(
            is_valid=False,
            error_message=f"Unsupported file type: {file.content_type}. Only PDF and DOCX files are supported.",
            file_size=file_size,
            file_type=file.content_type or "unknown"
        )
    
    # Validate file extension
    if file.filename:
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in SUPPORTED_TYPES[file.content_type]:
            return FileValidationResult(
                is_valid=False,
                error_message=f"File extension {file_extension} doesn't match content type {file.content_type}",
                file_size=file_size,
                file_type=file.content_type
            )
    
    return FileValidationResult(
        is_valid=True,
        file_size=file_size,
        file_type=file.content_type
    )

# Enhanced text extraction with better error handling
def extract_text_from_pdf(file_content: bytes) -> dict:
    """Enhanced PDF text extraction with metadata"""
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        text_content = ""
        page_count = len(doc)
        
        # Extract text from all pages
        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                text_content += f"--- Page {page_num + 1} ---\n{page_text}\n\n"
        
        doc.close()
        
        # Validate extracted text
        if not text_content.strip():
            raise ValueError("No readable text found in the PDF. The file might be image-based or corrupted.")
        
        # Calculate statistics
        word_count = len(text_content.split())
        character_count = len(text_content)
        
        return {
            "text": text_content.strip(),
            "page_count": page_count,
            "word_count": word_count,
            "character_count": character_count,
            "extraction_method": "PyMuPDF"
        }
        
    except Exception as e:
        if "No readable text found" in str(e):
            raise HTTPException(status_code=400, detail=str(e))
        else:
            raise HTTPException(status_code=400, detail=f"Error processing PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> dict:
    """Enhanced DOCX text extraction with metadata"""
    try:
        doc = Document(io.BytesIO(file_content))
        text_content = ""
        paragraph_count = 0
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
                paragraph_count += 1
        
        # Extract text from tables if any
        table_content = ""
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_content += " | ".join(row_text) + "\n"
        
        if table_content:
            text_content += "\n--- Tables ---\n" + table_content
        
        # Validate extracted text
        if not text_content.strip():
            raise ValueError("No readable text found in the DOCX file. The file might be empty or corrupted.")
        
        # Calculate statistics
        word_count = len(text_content.split())
        character_count = len(text_content)
        
        return {
            "text": text_content.strip(),
            "paragraph_count": paragraph_count,
            "table_count": len(doc.tables),
            "word_count": word_count,
            "character_count": character_count,
            "extraction_method": "python-docx"
        }
        
    except Exception as e:
        if "No readable text found" in str(e):
            raise HTTPException(status_code=400, detail=str(e))
        else:
            raise HTTPException(status_code=400, detail=f"Error processing DOCX: {str(e)}")

# Enhanced resume upload endpoint
@api_router.post("/resume/upload", response_model=FileUploadResponse)
async def upload_resume(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """Enhanced resume upload with comprehensive validation and processing"""
    
    # Validate file before processing
    validation_result = validate_uploaded_file(file)
    if not validation_result.is_valid:
        raise HTTPException(status_code=400, detail=validation_result.error_message)
    
    try:
        # Read file content
        file_content = await file.read()
        actual_file_size = len(file_content)
        
        # Validate file size after reading content
        MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
        if actual_file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400, 
                detail=f"File size ({actual_file_size / 1024 / 1024:.1f}MB) exceeds maximum allowed size (10MB)"
            )
        
        # Check for empty file
        if actual_file_size == 0:
            raise HTTPException(status_code=400, detail="Uploaded file is empty")
        
        # Extract text based on file type with enhanced processing
        extraction_result = None
        if file.content_type == "application/pdf":
            extraction_result = extract_text_from_pdf(file_content)
        else:  # DOCX
            extraction_result = extract_text_from_docx(file_content)
        
        extracted_text = extraction_result["text"]
        
        # Additional validation for resume content
        if len(extracted_text.split()) < 10:
            raise HTTPException(
                status_code=400, 
                detail="The extracted text is too short. Please ensure your resume contains sufficient content."
            )
        
        # Check for existing resume with same content
        existing_resume = await db.resumes.find_one({
            "user_id": current_user.id,
            "original_text": extracted_text
        })
        
        if existing_resume:
            raise HTTPException(
                status_code=409, 
                detail="A resume with identical content already exists. Please upload a different file."
            )
        
        # Create enhanced resume record
        resume = Resume(
            user_id=current_user.id,
            original_filename=file.filename,
            original_text=extracted_text
        )
        
        # Store in database with additional metadata
        resume_doc = resume.dict()
        resume_doc.update({
            "file_size": actual_file_size,
            "file_type": file.content_type,
            "extraction_metadata": extraction_result,
            "processing_status": "completed"
        })
        
        await db.resumes.insert_one(resume_doc)
        
        logger.info(f"Resume uploaded successfully - User: {current_user.email}, File: {file.filename}, Size: {actual_file_size} bytes")
        
        return FileUploadResponse(
            message="Resume uploaded and processed successfully",
            resume_id=resume.id,
            original_text=extracted_text,
            file_size=actual_file_size,
            word_count=extraction_result["word_count"],
            character_count=extraction_result["character_count"],
            file_type=file.content_type
        )
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        logger.error(f"Unexpected error during file upload - User: {current_user.email}, Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred while processing your file: {str(e)}")

@api_router.post("/resume/clean")
async def clean_resume(
    request: ResumeCleanRequest,
    current_user: User = Depends(get_current_user)
):
    # Get resume from database
    resume = await db.resumes.find_one({"id": request.resume_id, "user_id": current_user.id})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    try:
        # Initialize AI chat
        chat = await get_ai_chat()
        
        # Create user message for cleaning
        user_message = UserMessage(
            text=f"Please clean the following resume text by fixing all grammar and punctuation errors. Do not change the meaning, content structure, or formatting. Only fix grammatical errors, spelling mistakes, and punctuation. Return only the corrected text:\n\n{resume['original_text']}"
        )
        
        # Get AI response
        response = await chat.send_message(user_message)
        cleaned_text = response.strip()
        
        # Update resume with cleaned text
        await db.resumes.update_one(
            {"id": request.resume_id},
            {
                "$set": {
                    "cleaned_text": cleaned_text,
                    "updated_at": datetime.now(timezone.utc)
                }
            }
        )
        
        return {
            "message": "Resume cleaned successfully",
            "cleaned_text": cleaned_text
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error cleaning resume: {str(e)}")

@api_router.get("/resume/list")
async def list_resumes(current_user: User = Depends(get_current_user)):
    resumes = await db.resumes.find({"user_id": current_user.id}).to_list(100)
    return [Resume(**resume) for resume in resumes]

@api_router.get("/resume/{resume_id}")
async def get_resume(
    resume_id: str,
    current_user: User = Depends(get_current_user)
):
    resume = await db.resumes.find_one({"id": resume_id, "user_id": current_user.id})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    return Resume(**resume)

@api_router.get("/resume/{resume_id}/download")
async def download_resume(
    resume_id: str,
    current_user: User = Depends(get_current_user)
):
    resume = await db.resumes.find_one({"id": resume_id, "user_id": current_user.id})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    if not resume.get("cleaned_text"):
        raise HTTPException(status_code=400, detail="Resume has not been cleaned yet")
    
    try:
        # Create PDF from cleaned text
        filename = f"cleaned_{resume['original_filename'].rsplit('.', 1)[0]}"
        pdf_path = create_pdf_from_text(resume["cleaned_text"], filename)
        
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=f"{filename}.pdf"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()